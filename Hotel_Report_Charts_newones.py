import pandas as pd
from docx import Document
from docx.shared import Inches,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
import matplotlib.pyplot as plt
import math
from matplotlib.figure import Figure
import numpy as np
import os
import matplotlib
import locale
matplotlib.use('Agg')
locale.setlocale(locale.LC_ALL, 'it_IT')
matplotlib.style.use('seaborn')
df=pd.read_excel(r"C:\Users\sina.kian\Desktop\Ricardo\updates\DataBase_4_final.xlsx")
dz=pd.read_excel(r"C:\Users\sina.kian\Desktop\Ricardo\updates\Fina_Merge_LL.xlsx")
LL=dz[dz.columns[0]].to_list()
df=df[df.columns[1:]]

NN=['Fabbricato','Contenuto',
'Ricorso Terzi','Cristalli',
'Furto','Merci in Refrigerazione',
'Fenomeno Elettrico','Franchigia Danni Diretti',
'Margine di contribuzione annuo','Franchigia Danni Indiretti',
'Premio Imponibile Annuo Property 2021/2022','Premio Lordo Annuo Property 2021/2022',
'Fatturato  Hotel 2019 (€)', 'Fatturato Ristorante 2019 (€)',
'RC Terzi (RCT)','RC Prodotti (RCP) e Smercio',
'RC Prestatori di Lavoro (RCO)','RC Prestatori di Lavoro (RCO) - per persona','RC Terzi (RCT) Cose non consegnate / Garage keeper\'s liability',
'Franchigia (RCT)', 'Premio lordo Annuo 2019/2020', 'Fatturato  Hotel 2020 (€)', 'Fatturato  Hotel 2021 (€)', 'Fatturato  Hotel 2022 (€)']

dnum=df.copy()
for i in NN:
    for j in df[i].index:
        try:
            df[i][j]=locale.format_string('%10.2f', df[i][j], grouping=True)
        except:
            df[i][j]=df[i][j]
################################################################################

df['Codice_Unico']=range(0,len(df))
for i in range(0,len(df)):
    df['Codice_Unico'][i]=str(df['Codice Hotel'][i])+'_'+str(df['Denominazione Hotel'][i])

dnum['Codice_Unico']=range(0,len(dnum))
for i in range(0,len(dnum)):
    dnum['Codice_Unico'][i]=str(dnum['Codice Hotel'][i])+'_'+str(dnum['Denominazione Hotel'][i])

dff=df
#######
Anagrafica_Società=['Ragione Sociale', 'C.F./P.IVA', 'Sede Legale',
'Provincia (Sede Legale)', 'Città (Sede Legale)', 'Cap\n(Sede Legale)']
dT1= dff[Anagrafica_Società]
dT1.columns=['Ragione Sociale', 'C.F./P.IVA', 'Sede Legale',
'Provincia', 'Città', 'Cap']
##
##
Anagrafica_Hotel=['Denominazione Hotel', 'Codice Hotel', 'Indirizzo (Ubicazione Hotel)', 'Provincia \n(Ubicazione Hotel)',
       'Città (Ubicazione Hotel)', 'Cap (Ubicazione Hotel)', 'Assicurati Addizionali']
dT2= dff[Anagrafica_Hotel]
dT2.columns=['Denominazione Hotel', 'Codice Hotel', 'Indirizzo', 'Provincia',
       'Città', 'Cap', 'Assicurati Addizionali']
##
##
Danni_Diretti_T=['Zona rischi catastrofali Nr.', 'Presenza di vincolo', 'Fabbricato',
       'Contenuto', 'Ricorso Terzi', 'Cristalli', 'Furto',
       'Merci in Refrigerazione', 'Fenomeno Elettrico',
       'Terremoto, Inondazione, Alluvione, Allagamento','Terrorismo, Eventi Socio Politici, Atti Dolosi',
       'Franchigia Danni Diretti']

Danni_Diretti_N=['Zona rischi catastrofali Nr.', 'Presenza di vincolo', 'Fabbricato',
       'Contenuto', 'Ricorso Terzi', 'Cristalli', 'Furto',
       'Merci in Refrigerazione', 'Fenomeno Elettrico',
       'Coef. Terremoto',
       'Coef.  Inondazione, Alluvione, Allagamento','Coef. Terrorismo',
       'Coef. Eventi Socio Politici, Atti Dolosi',
       'Franchigia Danni Diretti']

dT3= dff[Danni_Diretti_T]
dN3=dnum[Danni_Diretti_N]

##
##
Danni_Indiretti=['Margine di contribuzione annuo',
       'Periodo di Indennizzo (Mesi)', 'Cimici da letto',
       'Cimici da letto\n Periodo di Indennizzo (n.mesi)',
       'Franchigia Danni Indiretti']

dT4=dff[Danni_Indiretti]
dN4=dnum[Danni_Indiretti]

##
##
Premio_Totale_PD_e_BI= ['Premio Imponibile Annuo Property 2021/2022',
       'Premio Lordo Annuo Property 2021/2022','Premio lordo Annuo 2019/2020']

dT5=dff[Premio_Totale_PD_e_BI]
##
##
Dati=['Fatturato  Hotel 2019 (€)', 'Fatturato  Hotel 2020 (€)','Fatturato  Hotel 2021 (€)','Fatturato  Hotel 2022 (€)','Fatturato Ristorante 2019 (€)','Franchigia (RCT)']
Massimali1= ['RC Terzi (RCT)',
       'RC Prodotti (RCP) e Smercio', 'RC Prestatori di Lavoro (RCO)',
       'RC Prestatori di Lavoro (RCO) - per persona',
       'RC Terzi (RCT) Cose non consegnate / Garage keeper\'s liability',
       'RC Terzi (RCT) Danni Patrimoniali',
       'RC Terzi (RCT) Danni Opere D\'Arte',
       'RC Terzi (RCT) Furto di beni consegnati',
       'RC Terzi (RCT) Annullamento Franchigia Furto']


dT61=dff[Dati]
dN61=dnum[Dati]

dT62=dff[Massimali1]
dN62=dnum[Massimali1]
##################################
################################
######################################
######################################
#######################################
for i in LL:
    #i=1
    WW = RGBColor(255, 255, 255)
    document = Document('temp.docx')
    style = document.styles['Normal']
    font = style.font
    font.name='Arial'
    #header = document.sections[0].header
    #paragraph = header.paragraphs[0]
    #logo_run = paragraph.add_run()
    #logo_run.add_picture("1_c_logo.png", width=Inches(1.75))
    #text_run = paragraph.add_run()
    #paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    #text_run.text = '\t' + '\t' + "© Copyright 2019 - Strategica Risk Consulting S.r.l. - Milano" # For center align of text
    #document.add_page_break()
    ##
    #T='{}, codice: {}'
    #document.add_heading(T.format(dff['Denominazione Hotel'][i],dff['Codice Hotel'][i]), 1)
    #
    #p=document.add_paragraph('Scheda Anagrafica (Società)')
    #p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ####################################################
    di1=dT1[dT1.index==i]
    di1=di1.T
    di1=di1.reset_index()
    di1.columns=['DATI SOCIETARI','Valori']
    ##
    table = document.add_table(rows=len(di1)+1, cols=3)
    table.style = 'Light Grid'
    ###################
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di1.columns[0]
    for ti in di1.index:
        table.cell(ti+1, 0).text=str(di1['DATI SOCIETARI'][ti])
        table.cell(ti+1, 1).text=str(di1['Valori'][ti])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    for nnn in np.arange(1,len(di1)+1):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    ######################################################
    di2=dT2[dT2.index==i]
    di2=di2.T
    di2=di2.reset_index()
    di2.columns=['DATI HOTEL','Valori']
    ##
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(1)
    table = document.add_table(rows=len(di2)+1, cols=3)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di2.columns[0]
    for ti in di2.index:
        table.cell(ti+1, 0).text=str(di2['DATI HOTEL'][ti])
        table.cell(ti+1, 1).text=str(di2['Valori'][ti])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    for nnn in np.arange(1,len(di2)+1):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    ###
    di3=dT3[dT3.index==i]
    di3=di3[['Fabbricato',
       'Contenuto', 'Ricorso Terzi', 'Cristalli', 'Furto',
       'Merci in Refrigerazione', 'Fenomeno Elettrico',
       'Terremoto, Inondazione, Alluvione, Allagamento',
       'Terrorismo, Eventi Socio Politici, Atti Dolosi','Franchigia Danni Diretti','Zona rischi catastrofali Nr.', 'Presenza di vincolo']]
    dn3=dN3[dN3.index==i]
    di3=di3.T
    di3=di3.reset_index()
    di3.columns=['SOMME ASSICURATE','VALORI ATTUALI']
    #_____________________
    dn31=dn3[dn3.columns[2:9]]
    dc31=dn3[dn3.columns[9:13]]
    x=['Fabbricato','Contenuto', 'Ricorso\n Terzi', 'Cristalli','Furto', 'Merci\n Refrigerazione', 'Fenomeno\n Elettrico']
    y=dn31.loc[i]
    xvals=[0,1.2,2.4,3,3.6,4.2,4.8]
    plt.figure(1)#,figsize=(10, 6))
    plt.yscale('log')
    mybars=plt.bar(xvals,(y),width=0.2,alpha=0.8)
    plt.tick_params(top='off', bottom='off', left='off', right='off', labelleft='off', labelbottom='on')
    x=plt.gca().xaxis
    for item in x.get_ticklabels():
        item.set_rotation(90)
    #
    for bari in mybars[:3]:
        height = bari.get_height()
        plt.gca().text(bari.get_x() + bari.get_width()/2,  0.15*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                     ha='center', color='w', fontsize=10, rotation=90, weight='bold')
    #
    for bari in mybars[3:]:
        height = bari.get_height()
        plt.gca().text(bari.get_x() + bari.get_width()/2,  1.1*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                     ha='center', color='k', fontsize=10, rotation=90, weight='bold')
    #
    new_xvals=[]
    for item in xvals[:2]:
        new_xvals.append(item+0.2)
    #
    mybars=plt.bar(new_xvals,dc31.loc[i][0]*(y[:2]),width=0.2,alpha=0.8)
    plt.tick_params(top='off', bottom='off', left='off', right='off', labelleft='off', labelbottom='on')
    for bari in mybars:
        height = bari.get_height()
        if height>0:
            plt.gca().text(bari.get_x() + bari.get_width()/2,  .2*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True), #11000
                         ha='center', color='w', fontsize=10, rotation=90, weight='bold')
        else:
            plt.gca().text(bari.get_x() + bari.get_width()/2,  11000, locale.format_string('%10.2f', (int(height)), grouping=True), #11000
                         ha='center', color='orange', fontsize=10, rotation=90, weight='bold')
    #
    new_xvals2=[]
    for item in new_xvals[:2]:
        new_xvals2.append(item+0.2)
    #
    mybars=plt.bar(new_xvals2,dc31.loc[i][1]*(y[:2]),width=0.2,alpha=0.8)
    plt.tick_params(top='off', bottom='off', left='off', right='off', labelleft='off', labelbottom='on')
    for bari in mybars:
        height = bari.get_height()
        if height>0:
            plt.gca().text(bari.get_x() + bari.get_width()/2, 0.2*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                         ha='center', color='w', fontsize=10, rotation=90, weight='bold')
        else:
            plt.gca().text(bari.get_x() + bari.get_width()/2,11000, locale.format_string('%10.2f', (int(height)), grouping=True),
                         ha='center', color='green', fontsize=10, rotation=90, weight='bold')
    #
    new_xvals3=[]
    for item in new_xvals2[:2]:
        new_xvals3.append(item+0.2)
    #
    mybars=plt.bar(new_xvals3,dc31.loc[i][2]*(y[:2]),width=0.2,alpha=0.8)
    plt.tick_params(top='off', bottom='off', left='off', right='off', labelleft='off', labelbottom='on')
    for bari in mybars:
        height = bari.get_height()
        if height>0:
            plt.gca().text(bari.get_x() + bari.get_width()/2, 0.2*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                     ha='center', color='w', fontsize=10, rotation=90, weight='bold')
        else:
            plt.gca().text(bari.get_x() + bari.get_width()/2, 11000, locale.format_string('%10.2f', (int(height)), grouping=True),
                     ha='center', color='red', fontsize=10, rotation=90, weight='bold')
    #
    new_xvals4=[]
    for item in new_xvals3[:2]:
        new_xvals4.append(item+0.2)
    #
    mybars=plt.bar(new_xvals4,dc31.loc[i][3]*(y[:2]),width=0.2,alpha=0.8)
    plt.tick_params(top='off', bottom='off', left='off', right='off', labelleft='off', labelbottom='on')
    for bari in mybars:
        height = bari.get_height()
        if height>0:
            plt.gca().text(bari.get_x() + bari.get_width()/2, 0.2*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                         ha='center', color='w', fontsize=10, rotation=90, weight='bold')
        else:
            plt.gca().text(bari.get_x() + bari.get_width()/2, 11000, locale.format_string('%10.2f', (int(height)), grouping=True),
                         ha='center', color='violet', fontsize=10, rotation=90, weight='bold')
    #
    fig = plt.gcf()
    fig.tight_layout()
    plt.tight_layout()
    ax=plt.gca()
    ax.set_xlabel('Partite Assicurate', weight='bold')
    ax.set_ylabel('Valori')
    ax.set_title('Rappresentazione Grafica Danni diretti')
    #axes = plt.axes()
    #axes.set_xticklabels(['fabbricato','contenuti', 'Ricorso Terzi', 'Cristalli','Furto', 'Merci Refrigerazione'])
    plt.xticks(xvals, ('Fabbricato','Contenuto', 'Ricorso\n Terzi', 'Cristalli','Furto', 'Merci\n Refrigerazione', 'Fenomeno\n Elettrico'))
    plt.legend(['Valori', 'Terremoto', 'Inondazione, Alluvione, Allagamento', 'Terrorismo', 'Eventi Socio Politici, Atti Dolosi' ],loc=1,frameon=False)
    for spine in plt.gca().spines.values():
        spine.set_visible(False)
    #
    plt.gca().axes.yaxis.set_visible(False)
    plt.gca().tick_params(axis='both', which='both', length=0)
    #matplotlib.style.use('default') ######################################################################################################################
    fname=df['Codice_Unico'][i]+'_plt1.png'
    try:
        fig.savefig(bbox_inches='tight',fname=fname)
    except:
        fig.savefig(fname=fname)
    #
    document.add_page_break()
    p=document.add_paragraph('Property Danni Diretti')
    table = document.add_table(rows=len(di3)+1, cols=3)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di3.columns[0]
    #hdr_cells[1].text = di3.columns[1]
    for ti in di3.index:
        table.cell(ti+1, 0).text=str(di3['SOMME ASSICURATE'][ti])
        table.cell(ti+1, 1).text=str(di3['VALORI ATTUALI'][ti])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    for nnn in np.arange(1,len(di3)+1):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    #--------
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(1)
    I=document.add_picture(fname, width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    os.remove(fname)
    plt.close('all')
    #document.add_paragraph('')
    #last_paragraph = document.paragraphs[-1]
    #last_paragraph.paragraph_format.space_before = Inches(1)
    ####
    di4=dT4[dT4.index==i]
    dn4=dN4[dN4.index==i]
    di4=di4.T
    di4=di4.reset_index()
    di4.columns=['SOMME ASSICURATE','VALORI ATTUALI']
    #_____________________
    x=[0,0,dn4['Periodo di Indennizzo (Mesi)'][i],dn4['Periodo di Indennizzo (Mesi)'][i]]
    y=[0,dn4['Margine di contribuzione annuo'][i],dn4['Margine di contribuzione annuo'][i],0]
    plt.figure(1)
    #matplotlib.style.use('seaborn')#########################################################################
    plt.plot(x,y,'o-',linewidth=3)
    #plt.xlim(left=-0.05)
    plt.ylim(bottom=-0)
    ax=plt.gca()
    ax.set_xlabel('Numero di Mesi')
    locs, labels = plt.xticks()  # Get the current locations and labels.
    plt.xticks(np.arange(0, dn4['Periodo di Indennizzo (Mesi)'][i]+6, step=6))  # Set label locations
    locs, labels = plt.yticks()  # Get the current locations and labels.
    plt.yticks(np.linspace(0, dn4['Margine di contribuzione annuo'][i],5))  # Set label locations
    ax.set_ylabel('Valori')
    ax.set_title('Rappresentazione Grafica\n Danni Indiretti (Margine di contribuzione)')
    fig = plt.gcf()
    fig.tight_layout()
    plt.tight_layout()
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(top=False,
                   bottom=True,
                   left=True,
                   right=False,
                   labelleft=True,
                   labelbottom=True)
    fname=df['Codice_Unico'][i]+'_plt2.png'
    fig.savefig(bbox_inches='tight',fname=fname)
    #
    #matplotlib.style.use('default')#########################################################################
    document.add_page_break()
    p=document.add_paragraph('Property Danni Indiretti')
    table = document.add_table(rows=len(di4)+1, cols=3)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di4.columns[0]
    for ti in di4.index:
        table.cell(ti+1, 0).text=str(di4['SOMME ASSICURATE'][ti])
        table.cell(ti+1, 1).text=str(di4['VALORI ATTUALI'][ti])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    for nnn in np.arange(1,len(di4)+1):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    #
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(1)
    I=document.add_picture(fname, width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    os.remove(fname)
    document.add_page_break()
    #
    plt.close('all')
    #####
    #####
    #document.add_page_break()
    di61=dT61[dT61.index==i]
    dn61=dN61[dN61.index==i]
    di61=di61.T
    di61=di61.reset_index()
    di61.columns=['DATI','VALORI ATTUALI']
    #
    di62=dT62[dT62.index==i]
    dn62=dN62[dN62.index==i]
    di62=di62.T
    di62=di62.reset_index()
    di62.columns=['DATI','VALORI ATTUALI']
    #_____________________
    #dn62=dn6[dn6.columns[2:7]]
    try:
        xx=dn61.loc[i][:-2].dropna().columns
    except:
        xx=dn61.loc[i][:-2].dropna().index
    #
    y1=dn61.loc[i][:-2].dropna()
    #y2=dn62.loc[i]
    #xvals=np.arange(len(xx))
    xvals=range(1,len(xx)+1)
    plt.figure(1)
    #plt.yscale('log')
    mybars=plt.bar(xvals,(y1),width=0.2,alpha=0.8)
    x=plt.gca().xaxis
    #for item in x.get_ticklabels():
    #    item.set_rotation(15)
    #
    Hh=[]
    for bari in mybars:
        Hh.append(bari.get_height())
        HM=np.max(Hh)
    #
    for bari in mybars:
        height = bari.get_height()
        if height>40000 and height>0.2*HM:
            plt.gca().text(bari.get_x() + bari.get_width()/2,  0.2*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                         ha='center', color='w', fontsize=10, rotation=90, weight='bold')
        else:
            plt.gca().text(bari.get_x() + bari.get_width()/2,  1.2*bari.get_height()+20000, locale.format_string('%10.2f', (int(height)), grouping=True),
                         ha='center', color='k', fontsize=10, rotation=90, weight='bold')
    #
    #mybars1=plt.bar(xvals+.1,(y2),width=0.2,alpha=0.8)
    #for item in x.get_ticklabels():
    #    item.set_rotation(15)
    #
    #for bari in mybars1:
    #    height = bari.get_height()
    #    if height>40000:
    #        plt.gca().text(bari.get_x() + bari.get_width()/2,  0.7*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
    #                     ha='center', color='w', fontsize=10, rotation=90, weight='bold')
    #    else:
    #        plt.gca().text(bari.get_x() + bari.get_width()/2,  0.2*bari.get_height()+20000, locale.format_string('%10.2f', (int(height)), grouping=True),
    #                     ha='center', color='k', fontsize=10, rotation=90, weight='bold')
    #
    fig = plt.gcf()
    fig.tight_layout()
    plt.tight_layout()
    ax=plt.gca()
    ax.set_xlabel('Fatturato', weight='bold')
    ax.set_ylabel('Valori')
    ax.set_title('Rappresentazione Grafica\n sull\'Andamento del Fatturato')
    #plt.xticks(xvals, ('2019', '2020'))
    lx=[]
    for kx in xx:
        lx.append(kx.split()[2])
    #
    plt.xticks(xvals, tuple(lx))
    #plt.legend(['Hotel', 'Ristorante'])
    for spine in plt.gca().spines.values():
        spine.set_visible(False)
    #
    plt.gca().axes.yaxis.set_visible(False)
    plt.gca().tick_params(axis='both', which='both', length=0)
    fname=df['Codice_Unico'][i]+'_plt3.png'
    try:
        fig.savefig(bbox_inches='tight',fname=fname)
    except:
        fig.savefig(fname=fname)
    #fname=df['Codice_Unico'][i]+'_plt3.png'
    #fig.savefig(bbox_inches='tight',fname=fname)
    ######
    document.add_paragraph('')
    p=document.add_paragraph('Liability')
    table = document.add_table(rows=len(di61)+1, cols=3)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di61.columns[0]
    #hdr_cells[1].text = di6.columns[1]
    for ti in di61.index:
        table.cell(ti+1, 0).text=str(di61['DATI'][ti])
        table.cell(ti+1, 1).text=str(di61['VALORI ATTUALI'][ti])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    for nnn in np.arange(1,len(di61)+1):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    #
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(0.6)
    I=document.add_picture(fname, width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    os.remove(fname)
    plt.close('all')
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(0.5)
    document.add_page_break()
    #
    #----------------------------------------------------------
    dn63=dn62[dn62.columns[0:5]]
    xx=['RC Terzi (RCT)', 'RC Prodotti (RCP) e Smercio',
           'RC Prestatori di Lavoro (RCO)',
           'RC Prestatori di Lavoro (RCO) - per persona',
           'RC Terzi (RCT) Cose non consegnate / Garage keeper\'s liability']
    #
    y=dn63.loc[i]
    xvals=np.arange(len(xx))
    plt.figure(1)
    #plt.yscale('log')
    mybars=plt.bar(xvals,y,width=0.2,color='red', alpha=0.8)
    x=plt.gca().xaxis
    for bari in mybars:
        height = bari.get_height()
        if height>300000:
            plt.gca().text(bari.get_x() + bari.get_width()/2,  0.6*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                         ha='center', color='w', fontsize=10, rotation=90, weight='bold')
        else:
            plt.gca().text(bari.get_x() + bari.get_width()/2,  1.1*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                         ha='center', color='k', fontsize=10, rotation=90, weight='bold')
    #for item in x.get_ticklabels():
    #    item.set_rotation(15)
    #
    fig = plt.gcf()
    fig.tight_layout()
    plt.tight_layout()
    ax=plt.gca()
    ax.set_xlabel('Responsabilità Civile', weight='bold')
    ax.set_ylabel('Valori')
    ax.set_title('Rapresentazione Grafica delle Somme Assicurate\n Responsabilità Civile')
    plt.xticks(xvals, ('RCT', 'RCP',
           'RCO',
           'RCO\n persona',
           'RCT\n non consegnate'))
    #
    for spine in plt.gca().spines.values():
        spine.set_visible(False)
    #
    plt.gca().axes.yaxis.set_visible(False)
    plt.gca().tick_params(axis='both', which='both', length=0)
    fname=df['Codice_Unico'][i]+'_plt4.png'
    fig.savefig(bbox_inches='tight',fname=fname)
    #
    document.add_paragraph('Primary')
    table = document.add_table(rows=len(di62)+1, cols=3)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ['Massimali']
    #hdr_cells[1].text = di6.columns[1]
    for ti in di62.index:
        table.cell(ti+1, 0).text=str(di62['DATI'][ti])
        table.cell(ti+1, 1).text=str(di62['VALORI ATTUALI'][ti])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    for nnn in np.arange(1,len(di62)+1):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    #
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(0.6)
    I=document.add_picture(fname, width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    plt.close('all')
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(0.5)
    #
    os.remove(fname)
    #######
    document.add_page_break()
    document.add_paragraph('Excess')
    table = document.add_table(rows=2, cols=3)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ['Massimali']
    #hdr_cells[1].text = di6.columns[1]
    table.cell(1, 0).text='RCT, RCP, RCO'
    table.cell(1, 1).text='18.000.000 xs 2.000.000'
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[1].height=Inches(0.2)
    ########################
    xx=['Hotel 1', 'Hotel 2', 'Hotel 3', 'Hotel 4', 'Hotel n']
    #
    y=[2000000]*len(xx)
    y2=[18000000]
    xvals=[0,0.3,0.6,0.9,1.2]
    xvals2=[.6]
    plt.figure(1)
    #plt.yscale('log')
    #matplotlib.style.use('seaborn')
    mybars=plt.bar(xvals,y,width=0.2) #fuchsia, ##18BA2B, #brown
    x=plt.gca().xaxis
    plt.gca().text(-.3 + bari.get_width()/2,  1.1*bari.get_height(), '2 Millioni',
                         ha='center', fontsize=10,rotation=90, weight='bold')
    #
    mybars=plt.bar(xvals2,y2,width=1.4,bottom=y)
    plt.gca().text(-.3 + bari.get_width()/2,  14000000, '18 Millioni',
                         ha='center', fontsize=10,rotation=90, weight='bold')
    #for item in x.get_ticklabels():
    #    item.set_rotation(15)
    #
    fig = plt.gcf()
    fig.tight_layout()
    plt.tight_layout()
    ax=plt.gca()
    #ax.set_xlabel('Responsabilità Civile', weight='bold')
    #ax.set_ylabel('Valori')
    ax.set_title('Rapresentazione Grafica delle Convenzione')
    plt.xticks(xvals, ('Hotel 1', 'Hotel 2', 'Hotel 3', 'Hotel 4', 'Hotel n'))
    plt.legend(['Massimali I Layer', 'Massimali II Layer' ],bbox_to_anchor=(1.04,1), loc="upper left",frameon=False)
    #plt.legend(bbox_to_anchor=(1.04,1), loc="upper left")
    #
    for spine in plt.gca().spines.values():
        spine.set_visible(False)
    #
    plt.gca().axes.yaxis.set_visible(False)
    plt.gca().tick_params(axis='both', which='both', length=0)
    fname=df['Codice_Unico'][i]+'_plt6.png'
    plt.savefig(bbox_inches='tight',fname=fname)
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(1)
    I=document.add_picture(fname, width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    os.remove(fname)
    plt.close('all')
    document.add_page_break()
    #######
    #document.add_paragraph('')
    table = document.add_table(rows=2, cols=3)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ['Franchigie']
    #hdr_cells[1].text = di6.columns[1]
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[1].height=Inches(0.2)
    #
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(1)
    xx=['Diretti', 'Indiretti', 'Liability']
    y=[dn3[dn3.columns[-1]][i],dn4[dn4.columns[-1]][i],dn61[dn61.columns[-1]][i]]
    xvals=np.arange(len(xx))
    plt.figure(1)
    #plt.yscale('log')
    mybars=plt.bar(xvals,y,width=0.2,color='darkviolet',alpha=0.8) #fuchsia, ##18BA2B, #brown
    x=plt.gca().xaxis
    Hh=[]
    for bari in mybars:
        Hh.append(bari.get_height())
        HM=np.max(Hh)
    #
    for bari in mybars:
        height = bari.get_height()
        if max(y)>3000:
            if height>3000:
                plt.gca().text(bari.get_x() + bari.get_width()/2,  0.75*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                             ha='center', color='w', fontsize=10, rotation=90, weight='bold')
            else:
                plt.gca().text(bari.get_x() + bari.get_width()/2,  1.1*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                             ha='center', color='k', fontsize=10, rotation=90, weight='bold')
        else:
            if height>300:
                plt.gca().text(bari.get_x() + bari.get_width()/2,  0.4*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                             ha='center', color='w', fontsize=10, rotation=90, weight='bold')
            else:
                plt.gca().text(bari.get_x() + bari.get_width()/2,  1.1*bari.get_height(), locale.format_string('%10.2f', (int(height)), grouping=True),
                             ha='center', color='k', fontsize=10, rotation=90, weight='bold')
    #
    fig = plt.gcf()
    fig.tight_layout()
    plt.tight_layout()
    ax=plt.gca()
    ax.set_xlabel('Franchigia', weight='bold')
    ax.set_ylabel('Valori')
    ax.set_title('Rapresentazione Grafica delle Franchigie Operanti')
    plt.xticks(xvals, ('Danni Diretti', 'Danni Indiretti', 'Liability'))
    for spine in plt.gca().spines.values():
        spine.set_visible(False)
    #
    plt.gca().axes.yaxis.set_visible(False)
    plt.gca().tick_params(axis='both', which='both', length=0)
    fname=df['Codice_Unico'][i]+'_plt5.png'
    fig.savefig(bbox_inches='tight',fname=fname)
    I=document.add_picture(fname, width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    os.remove(fname)
    plt.close('all')
    #############
    di5=dT5[dT5.index==i]
    di5=di5.T
    di5=di5.reset_index()
    di5.columns=['PREMIO','VALORI ATTUALI']
    document.add_page_break()
    p=document.add_paragraph('Premio Totale PD + BI')
    table = document.add_table(rows=len(di5)+1, cols=3)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di5.columns[0]
    #hdr_cells[1].text = di5.columns[1]
    for ti in di5.index:
        table.cell(ti+1, 0).text=str(di5['PREMIO'][ti])
        table.cell(ti+1, 1).text=str(di5['VALORI ATTUALI'][ti])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    for nnn in np.arange(1,len(di5)+1):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    ########
    #di7=dT7[dT7['Codice_Unico']==dT7['Codice_Unico'][i]]
    #di7=di7.T
    #di7.columns=['jshjs']
    #di7=di7.reset_index()
    #di7.columns=['DATI','VALORI ATTUALI']
    #######
    #document.add_paragraph('')
    #last_paragraph = document.paragraphs[-1]
    #last_paragraph.paragraph_format.space_before = Inches(1)
    #p=document.add_paragraph('Dati Amministrativi')
    #table = document.add_table(rows=len(di7)-1, cols=3)
    #table.style = 'Light Grid'
    #hdr_cells = table.rows[0].cells
    #hdr_cells[0].text = di7.columns[0]
    #hdr_cells[1].text = di7.columns[1]
    #for ti in di7.index[1:-1]:
    #    table.cell(ti, 0).text=str(di7['DATI'][ti])
    #    table.cell(ti, 1).text=str(di7['VALORI ATTUALI'][ti])
    #table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    #table.rows[0].height=Inches(0.3)
    #for nnn in np.arange(1,len(di7)-1):
    #    table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    #    table.rows[nnn].height=Inches(0.2)
    ###################
    ####################
    name=df['Codice_Unico'][i]+'.docx'
    document.save(name)
    plt.close('all')
