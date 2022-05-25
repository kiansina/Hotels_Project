import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
import matplotlib.pyplot as plt
import math
from matplotlib.figure import Figure

df=pd.read_excel(r"*****")
df.columns=df.loc[5]
df=df[6:168].set_index(i for i in range(0,len(df[6:168])))
df=df.loc[:,~df.columns.duplicated()]
df['Franchigia Danni Indiretti']=['5 Giorni']*len(df)
df['Fatturato totale']=[0]*len(df)
for i in range(0,len(df)):
    df['Fatturato totale'][i]=df['Fatturato  HOTEL 2019'][i]+df['Fatturato Ristorante 2019'][i]

NN=['Franchigia Danni Diretti',
       'Valore fabbricato', 'Valore contenuti', 'Ricorso Terzi', 'Cristalli',
       'Furto', 'Merci Refrigerazione', 'Fenomeno Elettrico', 'Margine di contribuzione', 'TOTALE FABBRICATO + CONTENUTO', 'Fatturato Ristorante 2019',
       'Fatturato Ristorante 2020', 'Fatturato  HOTEL 2019',
       'Fatturato HOTEL 2020','RC Terzi (RCT)', 'RC Prodotti (RCP) e Smercio',
       'RC Prestatori di Lavoro (RCO)',
       'RC Prestatori di Lavoro (RCO) - per persona']

#NN=['Franchigia Danni Diretti',
#       'Valore fabbricato', 'Valore contenuti', 'Ricorso Terzi', 'Cristalli',
#       'Furto', 'Merci Refrigerazione', 'Fenomeno Elettrico', 'Margine di contribuzione', 'TOTALE FABBRICATO + CONTENUTO', 'Premio Imponibile Annuo',
#       'Premio Imponibile PRORATA', 'Premio Lordo Annuo',
#       'Premio Lordo PRORATA', 'Franchigia (RCT)', 'Fatturato Ristorante 2019',
#       'Fatturato Ristorante 2020', 'Fatturato  HOTEL 2019',
#       'Fatturato HOTEL 2020', 'RC Terzi (RCT)', 'RC Prodotti (RCP) e Smercio',
#       'RC Prestatori di Lavoro (RCO)',
#       'RC Prestatori di Lavoro (RCO) - per persona','RC Terzi (RCT) Cose non consegnate / Garage keeper\'s liability',
#       'Premio lordo Annuo 2019', 'Premio Lordo Annuo 2020', 'Fatturato totale']
for i in NN:
    df[i].fillna(0,inplace=True)

dnum=df.copy()
for i in NN:
        df[i]=df[i].map('{:,d}'.format)


df['Codice_Unico']=range(0,len(df))
for i in range(0,len(df)):
    df['Codice_Unico'][i]=str(df['Codice Hotel'][i])+'_'+str(df['Denominazione Hotel'][i])

dnum['Codice_Unico']=range(0,len(dnum))
for i in range(0,len(dnum)):
    dnum['Codice_Unico'][i]=str(dnum['Codice Hotel'][i])+'_'+str(dnum['Denominazione Hotel'][i])

dff=df[:8]
#######
Anagrafica_Società=['Codice Hotel', 'Numero Certificato', 'Assicurato',
       'C.F./P.IVA', 'Indirizzo (Sede Legale)', 'Provincia (Sede Legale)',
       'Città (Sede Legale)', 'Cap\n(Sede Legale)','Codice_Unico']

dT1= dff[Anagrafica_Società]
##
##
Anagrafica_Hotel=['Codice Hotel', 'Indirizzo (Ubicazione Hotel)', 'Provincia \n(Ubicazione Hotel)',
       'Città (Ubicazione Hotel)', 'Cap (Ubicazione Hotel)',
       'Denominazione Hotel','Codice_Unico']

dT2= dff[Anagrafica_Hotel]
##
##
Danni_Diretti=['Codice Hotel', 'Zona rischi catastrofali Nr.', 'Valore fabbricato', 'Valore contenuti', 'Ricorso Terzi', 'Cristalli',
       'Furto', 'Merci Refrigerazione', 'Fenomeno Elettrico', 'Terremoto, Inondazione, Alluvione, Allagamento',
       'Terrorismo, Eventi Socio Politici, Atti Dolosi', 'Franchigia Danni Diretti','Codice_Unico']


dT3= dff[Danni_Diretti]
dN3=dnum[Danni_Diretti]

##
##
Danni_Indiretti=['Codice Hotel', 'Margine di contribuzione', 'Periodo di Indennizzo (Mesi)', 'Cimici da letto\n(se 0 o "-" non operante; numero mesi se operante)','Franchigia Danni Indiretti','Codice_Unico']

dT4=dff[Danni_Indiretti]
dN4=dnum[Danni_Indiretti]
dT4.columns=['Codice Hotel', 'Margine di contribuzione Annuo',
       'Periodo di Indennizzo (Mesi)',
       'Cimici da letto\n(se 0 o "-" non operante; numero mesi se operante)','Franchigia Danni Indiretti',
       'Codice_Unico']
##
##
Premio_Totale_PD_e_BI= ['Codice Hotel', 'Premio Imponibile Annuo', 'Premio Imponibile PRORATA', 'Premio Lordo Annuo', 'Premio Lordo PRORATA','Codice_Unico']

dT5=dff[Premio_Totale_PD_e_BI]
##
##
Sezione_Liability= ['Codice Hotel', 'Fatturato  HOTEL 2019', 'Fatturato HOTEL 2020', 'Fatturato Ristorante 2019',
       'Fatturato Ristorante 2020', 'RC Terzi (RCT)', 'RC Prodotti (RCP) e Smercio',
       'RC Prestatori di Lavoro (RCO)',
       'RC Prestatori di Lavoro (RCO) - per persona',
       'RC Terzi (RCT) Danni Patrimoniali',
       'RC Terzi (RCT) Danni Opere D\'Arte', 'RC Terzi (RCT) Furto di beni consegnati',
       'RC Terzi (RCT) Annullamento Franchigia Furto',
       'RC Terzi (RCT) Cose non consegnate / Garage keeper\'s liability', 'Franchigia (RCT)', 'Premio lordo Annuo 2019', 'Premio Lordo Annuo 2020','Codice_Unico']


dT6=dff[Sezione_Liability]
dN6=dnum[Sezione_Liability]
##
##
Dati_Amministratitivi=['Codice Hotel', 'Effetto della copertura\n H 24.00 del', 'Scadenza della copertura\n H 24.00 del', 'Presenza di vincolo', 'Appendici Property ', 'Appendici Casualty',
       'Automatico / Manuale','Codice_Unico']

dT7=dff[Dati_Amministratitivi]
##
##
for i in range(0,len(dff['Codice Hotel'])):
    WW = RGBColor(255, 255, 255)
    document = Document('temp.docx')
    style = document.styles['Normal']
    font = style.font
    font.name='Arial'
    #header = document.sections[0].header
    #paragraph = header.paragraphs[0]
    #logo_run = paragraph.add_run()
    #logo_run.add_picture("1_c_logo.png", width=Inches(1.75))
    #text_run = paragraph.add_run()
    #paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    #text_run.text = '\t' + '\t' + "© Copyright 2019 - Strategica Risk Consulting S.r.l. - Milano" # For center align of text
    document.add_page_break()
    ##
    T='{}, codice: {}'
    document.add_heading(T.format(dff['Denominazione Hotel'][i],dff['Codice Hotel'][i]), 1)
    #
    p=document.add_paragraph('Scheda Anagrafica (Società)')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ####################################################
    di1=dT1[dT1['Codice_Unico']==dT1['Codice_Unico'][i]]
    di1=di1.T
    di1.columns=['jshjs'] #When I don't put it in one combination column name becomes 5 and also index becomes 5 and problem happens
    di1=di1.reset_index()
    di1.columns=['DATI SOCIETARI','ATTUALI']
    ##
    table = document.add_table(rows=len(di1)-1, cols=2)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di1.columns[0]
    hdr_cells[1].text = di1.columns[1]
    for ti in di1.index[1:-1]:
        table.cell(ti, 0).text=str(di1['DATI SOCIETARI'][ti])
        table.cell(ti, 1).text=str(di1['ATTUALI'][ti])
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    ######################################################
    di2=dT2[dT2['Codice_Unico']==dT2['Codice_Unico'][i]]
    di2=di2.T
    di2.columns=['jshjs']
    di2=di2.reset_index()
    di2.columns=['DATI HOTEL','ATTUALI']
    ##
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(1)
    p=document.add_paragraph('Scheda Anagrafica (Hotel)')
    table = document.add_table(rows=len(di2), cols=2)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di2.columns[0]
    hdr_cells[1].text = di2.columns[1]
    for ti in di2.index[:-1]:
        table.cell(ti+1, 0).text=str(di2['DATI HOTEL'][ti])
        table.cell(ti+1, 1).text=str(di2['ATTUALI'][ti])
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    ###
    di3=dT3[dT3['Codice_Unico']==dT3['Codice_Unico'][i]]
    dn3=dN3[dN3['Codice_Unico']==dN3['Codice_Unico'][i]]
    di3=di3.T
    di3.columns=['jshjs']
    di3=di3.reset_index()
    di3.columns=['PARTITE ASSICURATE','VALORI ATTUALI']
    #_____________________
    dn31=dn3[dn3.columns[2:8]]
    x=['fabbricato','contenuti', 'Ricorso Terzi', 'Cristalli','Furto', 'Merci Refrigerazione']
    y=dn31.loc[i]
    xvals=range(len(x))
    plt.figure()
    plt.yscale('log')
    plt.bar(xvals,(y),width=0.2)
    x=plt.gca().xaxis
    for item in x.get_ticklabels():
        item.set_rotation(15)
    #
    new_xvals=[]
    for item in xvals[:2]:
        new_xvals.append(item+0.2)
    #
    plt.bar(new_xvals,0.5*(y[:2]),width=0.2,color='red')
    #
    new_xvals2=[]
    for item in new_xvals[:2]:
        new_xvals2.append(item+0.2)
    #
    plt.bar(new_xvals2,0.7*(y[:2]),width=0.2,color='yellow')
    #
    fig = plt.gcf()
    fig.tight_layout()
    plt.tight_layout()
    ax=plt.gca()
    ax.set_xlabel('Partite Assicurate')
    ax.set_ylabel('Valori')
    ax.set_title('Danni diretti')
    #axes = plt.axes()
    #axes.set_xticklabels(['fabbricato','contenuti', 'Ricorso Terzi', 'Cristalli','Furto', 'Merci Refrigerazione'])
    plt.xticks(xvals, ('fabbricato','contenuti', 'Ricorso Terzi', 'Cristalli','Furto', 'Merci Refrigerazione'))
    plt.legend(['Valori', 'Per Terremoto', 'per terrorismo'])
    fname=df['Codice_Unico'][i]+'_plt1'
    fig.savefig(bbox_inches='tight',fname=fname)
    ###
    document.add_page_break()
    p=document.add_paragraph('Danni Diretti')
    table = document.add_table(rows=len(di3)-1, cols=2)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di3.columns[0]
    hdr_cells[1].text = di3.columns[1]
    for ti in di3.index[1:-1]:
        table.cell(ti, 0).text=str(di3['PARTITE ASSICURATE'][ti])
        table.cell(ti, 1).text=str(di3['VALORI ATTUALI'][ti])
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    #--------
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(1)
    I=document.add_picture(fname+'.png', width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    ####
    di4=dT4[dT4['Codice_Unico']==dT4['Codice_Unico'][i]]
    dn4=dN4[dN4['Codice_Unico']==dN4['Codice_Unico'][i]]
    di4=di4.T
    di4.columns=['jshjs']
    di4=di4.reset_index()
    di4.columns=['PARTITE ASSICURATE','VALORI ATTUALI']
    #_____________________
    x=[0,0,dn4['Periodo di Indennizzo (Mesi)'][i],dn4['Periodo di Indennizzo (Mesi)'][i]]
    y=[0,dn4['Margine di contribuzione'][i],dn4['Margine di contribuzione'][i],0]
    plt.figure()
    plt.plot(x,y,'-')
    #plt.xlim(left=-0.05)
    plt.ylim(bottom=-0)
    ax=plt.gca()
    ax.set_xlabel('numero di mesi')
    ax.set_ylabel('Valori')
    ax.set_title('Danni Indiretti (Margine di contribuzione)')
    fig = plt.gcf()
    fig.tight_layout()
    plt.tight_layout()
    fname=df['Codice_Unico'][i]+'_plt2'
    fig.savefig(bbox_inches='tight',fname=fname)
    ###
    p=document.add_paragraph('Danni Indiretti')
    table = document.add_table(rows=len(di4)-1, cols=2)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di4.columns[0]
    hdr_cells[1].text = di4.columns[1]
    for ti in di4.index[1:-1]:
        table.cell(ti, 0).text=str(di4['PARTITE ASSICURATE'][ti])
        table.cell(ti, 1).text=str(di4['VALORI ATTUALI'][ti])
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    #--------
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(1)
    I=document.add_picture(fname+'.png', width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #####
    di5=dT5[dT5['Codice_Unico']==dT5['Codice_Unico'][i]]
    di5=di5.T
    di5.columns=['jshjs']
    di5=di5.reset_index()
    di5.columns=['PREMIO','VALORI ATTUALI']
    #####
    document.add_page_break()
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(.5)
    p=document.add_paragraph('Premio Totale PD + BI')
    table = document.add_table(rows=len(di5)-1, cols=2)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di5.columns[0]
    hdr_cells[1].text = di5.columns[1]
    for ti in di5.index[1:-1]:
        table.cell(ti, 0).text=str(di5['PREMIO'][ti])
        table.cell(ti, 1).text=str(di5['VALORI ATTUALI'][ti])
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    ######
    di6=dT6[dT6['Codice_Unico']==dT6['Codice_Unico'][i]]
    dn6=dN6[dN6['Codice_Unico']==dN6['Codice_Unico'][i]]
    di6=di6.T
    di6.columns=['jshjs']
    di6=di6.reset_index()
    di6.columns=['DATI','VALORI ATTUALI']
    #_____________________
    dn61=dn6[dn6.columns[1:3]]
    dn62=dn6[dn6.columns[3:5]]
    xx=['2019','2020']
    y1=dn61.loc[i]
    y2=dn62.loc[i]
    xvals=np.arange(len(xx))
    plt.figure()
    #plt.yscale('log')
    plt.bar(xvals-.1,(y1),width=0.2)
    x=plt.gca().xaxis
    for item in x.get_ticklabels():
        item.set_rotation(15)
    #
    plt.bar(xvals+.1,(y2),width=0.2,color='red')
    fig = plt.gcf()
    fig.tight_layout()
    plt.tight_layout()
    ax=plt.gca()
    ax.set_xlabel('Fatturato')
    ax.set_ylabel('Valori')
    ax.set_title('Liability-Fatturati')
    plt.xticks(xvals, ('2019', '2020'))
    plt.legend(['Hotel', 'Ristorante'])
    fname=df['Codice_Unico'][i]+'_plt3'
    fig.savefig(bbox_inches='tight',fname=fname)
    ######
    #--------
    document.add_paragraph('')
    p=document.add_paragraph('Sezione Liability')
    table = document.add_table(rows=len(di6)-1, cols=2)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di6.columns[0]
    hdr_cells[1].text = di6.columns[1]
    for ti in di6.index[1:-1]:
        table.cell(ti, 0).text=str(di6['DATI'][ti])
        table.cell(ti, 1).text=str(di6['VALORI ATTUALI'][ti])
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    #_#_
    document.add_paragraph('')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_before = Inches(0.6)
    I=document.add_picture(fname+'.png', width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    #----------------------------------------------------------
    dn63=dn6[dn6.columns[[5,6,7,8,13]]]
    xx=['RC Terzi (RCT)', 'RC Prodotti (RCP) e Smercio',
           'RC Prestatori di Lavoro (RCO)',
           'RC Prestatori di Lavoro (RCO) - per persona',
           'RC Terzi (RCT) Cose non consegnate / Garage keeper\'s liability']
    #
    y=dn63.loc[i]
    xvals=np.arange(len(xx))
    plt.figure()
    #plt.yscale('log')
    plt.bar(xvals,y,width=0.2)
    x=plt.gca().xaxis
    for item in x.get_ticklabels():
        item.set_rotation(15)
    #
    fig = plt.gcf()
    fig.tight_layout()
    plt.tight_layout()
    ax=plt.gca()
    ax.set_xlabel('Risponibilità Civile')
    ax.set_ylabel('Valori')
    ax.set_title('Liability - Risponibilità Civile')
    plt.xticks(xvals, ('RC Terzi (RCT)', 'RC Prodotti (RCP) e Smercio',
           'RC Prestatori di Lavoro (RCO)',
           'RC Prestatori di Lavoro (RCO) - per persona',
           'RC Terzi (RCT) Cose non consegnate / Garage keeper\'s liability'))
    fname=df['Codice_Unico'][i]+'_plt4'
    fig.savefig(bbox_inches='tight',fname=fname)
    I=document.add_picture(fname+'.png', width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #######
    document.add_paragraph('')
    xx=['Diretti', 'Indiretti', 'Liability']
    y=[dn3[dn3.columns[-2]][i],dn4[dn4.columns[1]][i]*5/365,dn6[dn6.columns[-4]][i]]
    xvals=np.arange(len(xx))
    plt.figure()
    #plt.yscale('log')
    plt.bar(xvals,y,width=0.2)
    x=plt.gca().xaxis
    for item in x.get_ticklabels():
        item.set_rotation(15)
    #
    fig = plt.gcf()
    fig.tight_layout()
    plt.tight_layout()
    ax=plt.gca()
    ax.set_xlabel('Franchigia')
    ax.set_ylabel('Valori')
    ax.set_title('Franchigia di Danni')
    plt.xticks(xvals, ('Diretti', 'Indiretti', 'Liability'))
    fname=df['Codice_Unico'][i]+'_plt5'
    fig.savefig(bbox_inches='tight',fname=fname)
    I=document.add_picture(fname+'.png', width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ########
    di7=dT7[dT7['Codice_Unico']==dT7['Codice_Unico'][i]]
    di7=di7.T
    di7.columns=['jshjs']
    di7=di7.reset_index()
    di7.columns=['DATI','VALORI ATTUALI']
    #######
    document.add_page_break()
    p=document.add_paragraph('Dati Amministrativi')
    table = document.add_table(rows=len(di7)-1, cols=2)
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = di7.columns[0]
    hdr_cells[1].text = di7.columns[1]
    for ti in di7.index[1:-1]:
        table.cell(ti, 0).text=str(di7['DATI'][ti])
        table.cell(ti, 1).text=str(di7['VALORI ATTUALI'][ti])
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.3)
    name=df['Codice_Unico'][i]+'.docx'
    document.save(name)
