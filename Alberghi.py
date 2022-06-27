import pandas as pd
from docx import Document
from docx.shared import Inches,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
import matplotlib.pyplot as plt
import math
from matplotlib.figure import Figure
import numpy as np
import os
import matplotlib
import locale
matplotlib.use('Agg')
locale.setlocale(locale.LC_ALL, 'it_IT')
matplotlib.style.use('seaborn')
df=pd.read_excel(r"C:\Users\sina.kian\Desktop\Ricardo\updates\DataBase_5_final.xlsx")





df=df[df.columns[1:]]

NN=['Fabbricato','Contenuto',
'Ricorso Terzi','Cristalli',
'Furto','Merci in Refrigerazione',
'Fenomeno Elettrico','Franchigia Danni Diretti',
'Margine di contribuzione annuo','Franchigia Danni Indiretti',
'Premio Imponibile Annuo Property 2021/2022','Premio Lordo Annuo Property 2021/2022',
'Fatturato  Hotel 2019 (€)', 'Fatturato Ristorante 2019 (€)',
'RC Terzi (RCT)','RC Prodotti (RCP) e Smercio',
'RC Prestatori di Lavoro (RCO)','RC Prestatori di Lavoro (RCO) - per persona','RC Terzi (RCT) Cose non consegnate / Garage keeper\'s liability',
'Franchigia (RCT)', 'Premio lordo Annuo 2019/2020', 'Fatturato  Hotel 2020 (€)', 'Fatturato  Hotel 2021 (€)', 'Fatturato  Hotel 2022 (€)']

dnum=df.copy()
for i in NN:
    for j in df[i].index:
        try:
            df[i][j]=locale.format_string('%10.0f €', df[i][j], grouping=True)
        except:
            df[i][j]=df[i][j]


df['Codice_Unico']=range(0,len(df))
for i in range(0,len(df)):
    df['Codice_Unico'][i]=str(df['Codice Hotel'][i])+'_'+str(df['Denominazione Hotel'][i])

dnum['Codice_Unico']=range(0,len(dnum))
for i in range(0,len(dnum)):
    dnum['Codice_Unico'][i]=str(dnum['Codice Hotel'][i])+'_'+str(dnum['Denominazione Hotel'][i])

dff=df
########################
##########################
########################
#########################
#######################
#########################

Anagrafica_Società=['Ragione Sociale', 'C.F./P.IVA', 'Sede Legale',
'Provincia (Sede Legale)', 'Città (Sede Legale)', 'Cap\n(Sede Legale)']
dT1= dff[Anagrafica_Società]
dT1.columns=['Ragione Sociale', 'C.F./P.IVA', 'Sede Legale',
'Provincia', 'Città', 'Cap']
##
##
Anagrafica_Hotel=['Denominazione Hotel', 'Codice Hotel', 'Indirizzo (Ubicazione Hotel)', 'Provincia \n(Ubicazione Hotel)',
       'Città (Ubicazione Hotel)', 'Cap (Ubicazione Hotel)', 'Assicurati Addizionali']
dT2= dff[Anagrafica_Hotel]
dT2.columns=['Denominazione Hotel', 'Codice Hotel', 'Indirizzo', 'Provincia',
       'Città', 'Cap', 'Assicurati Addizionali']
##
##
Danni_Diretti_T=['Zona rischi catastrofali Nr.', 'Presenza di vincolo', 'Fabbricato',
       'Contenuto', 'Ricorso Terzi', 'Cristalli', 'Furto',
       'Merci in Refrigerazione', 'Fenomeno Elettrico',
       'Terremoto, Inondazione, Alluvione, Allagamento','Terrorismo, Eventi Socio Politici, Atti Dolosi',
       'Franchigia Danni Diretti','Coef. Terremoto',
        'Coef.  Inondazione, Alluvione, Allagamento']

Danni_Diretti_N=['Zona rischi catastrofali Nr.', 'Presenza di vincolo', 'Fabbricato',
       'Contenuto', 'Ricorso Terzi', 'Cristalli', 'Furto',
       'Merci in Refrigerazione', 'Fenomeno Elettrico',
       'Coef. Terremoto',
       'Coef.  Inondazione, Alluvione, Allagamento','Coef. Terrorismo',
       'Coef. Eventi Socio Politici, Atti Dolosi',
       'Franchigia Danni Diretti']

dT3= dff[Danni_Diretti_T]
dT3.columns=['Zona rischi catastrofali Nr.', 'Presenza di vincolo', 'Fabbricato',
       'Contenuto', 'Ricorso Terzi', 'Cristalli', 'Furto',
       'Merci in Refrigerazione', 'Fenomeno Elettrico',
       'Terremoto, Inondazione, Alluvione, Allagamento','Terrorismo, Eventi Socio Politici, Atti Dolosi',
       'Franchigia Danni Diretti','Limite di Indennizzo (S.A) Terremoto',
        'Limite di Indennizzo (S.A)  Inondazione, Alluvione, Allagamento']
dN3=dnum[Danni_Diretti_N]

##
##
Danni_Indiretti=['Margine di contribuzione annuo',
       'Periodo di Indennizzo (Mesi)', 'Cimici da letto',
       'Cimici da letto\n Periodo di Indennizzo (n.mesi)',
       'Franchigia Danni Indiretti']

dT4=dff[Danni_Indiretti]
dT4.columns=['Margine di contribuzione annuo',
       'Periodo di Indennizzo (Mesi)', 'Cimici da letto',
       'Periodo di Indennizzo (Mesi)',
       'Franchigia Danni Indiretti (5 Giorni)']
dN4=dnum[Danni_Indiretti]

##
##
Premio_Totale_PD_e_BI= ['Premio Imponibile Annuo Property 2021/2022',
       'Premio Lordo Annuo Property 2021/2022','Premio lordo Annuo 2019/2020']

dT5=dff[Premio_Totale_PD_e_BI]
##
##
Dati=['Fatturato  Hotel 2019 (€)', 'Fatturato  Hotel 2020 (€)','Fatturato  Hotel 2021 (€)','Fatturato  Hotel 2022 (€)','Fatturato Ristorante 2019 (€)','Franchigia (RCT)']
Massimali1= ['RC Terzi (RCT)',
       'RC Prodotti (RCP) e Smercio', 'RC Prestatori di Lavoro (RCO)',
       'RC Prestatori di Lavoro (RCO) - per persona',
       'RC Terzi (RCT) Cose non consegnate / Garage keeper\'s liability',
       'RC Terzi (RCT) Danni Patrimoniali',
       'RC Terzi (RCT) Danni Opere D\'Arte',
       'RC Terzi (RCT) Furto di beni consegnati',
       'RC Terzi (RCT) Annullamento Franchigia Furto']


dT61=dff[Dati]
dN61=dnum[Dati]

dT62=dff[Massimali1]
dN62=dnum[Massimali1]
########################
##########################
########################
#########################
#######################
#########################
for i in range(0,len(dff['Codice Hotel'])):
    #i=1
    WW = RGBColor(255, 255, 255)
    document = Document('temp.docx')
    style = document.styles['Normal']
    font = style.font
    font.name='Times New Roman'
    ##
    ##
    ##
    ##
    SSS=len('Assicurato: ')+len(dff['Denominazione Hotel'][i])+len('Contraente: ')+len('Best Western Italia Scpa')
    p=document.add_paragraph('');
    p.add_run('Assicurato: ').bold=True;
    if SSS<53:
        NS=188
    elif SSS<58:
        NS=183
    elif SSS<63:
        NS=178
    elif SSS<68:
        NS=173
    elif SSS<73:
        NS=168
    elif SSS<78:
        NS=163
    elif SSS<83:
        NS=158
    p.add_run('{}'.format(dff['Denominazione Hotel'][i])+' '*(NS-SSS));
    p.add_run('Contraente: ').bold=True;
    p.add_run('Best Western Italia Scpa');
    SSS=len('Codice Fiscale/P.IVA: ')+len(str(dff['C.F./P.IVA'][i]))+len(str(dff['Sede Legale'][i])+'-'+str(dff['Città (Sede Legale)'][i])+'-'+str(dff['Cap\n(Sede Legale)'][i])+'- ('+str(dff['Provincia (Sede Legale)'][i])+')')+len('Sede legale: ')
    p.add_run('\nCodice Fiscale/P.IVA: ').bold=True;
    if SSS<53:
        NS=188
    elif SSS<58:
        NS=183
    elif SSS<63:
        NS=178
    elif SSS<68:
        NS=173
    elif SSS<73:
        NS=168
    elif SSS<78:
        NS=163
    elif SSS<83:
        NS=158
    elif SSS<88:
        NS=153
    elif SSS<93:
        NS=148
    elif SSS<98:
        NS=143
    p.add_run('{}'.format(dff['C.F./P.IVA'][i])+' '*(NS-SSS));
    p.add_run('Sede legale: ').bold=True;
    p.add_run('{}'.format(str(dff['Sede Legale'][i]).capitalize()+'-'+str(dff['Città (Sede Legale)'][i])+'-'+str(dff['Cap\n(Sede Legale)'][i])+'- ('+str(dff['Provincia (Sede Legale)'][i])+')'))
    p=document.add_paragraph('')#
    p.add_run('Hotel Assicurato').bold=True
    SSS=len('Denominazione Hotel: ')+len('Property Code: ')+len(dff['Denominazione Hotel'][i])+len(str(dff['Codice Hotel'][i]))
    if SSS<53:
        NS=185
    elif SSS<58:
        NS=180
    elif SSS<63:
        NS=175
    elif SSS<68:
        NS=170
    elif SSS<73:
        NS=165
    elif SSS<78:
        NS=160
    elif SSS<83:
        NS=155
    elif SSS<88:
        NS=150
    elif SSS<93:
        NS=145
    elif SSS<98:
        NS=140
    p.add_run('\nDenominazione Hotel: {}'.format(dff['Denominazione Hotel'][i])+' '*(NS-SSS)+ 'Property Code: {}'.format(dff['Codice Hotel'][i]))
    SSS=len('Ubicazione Hotel: ')+len(str(dff['Indirizzo (Ubicazione Hotel)'][i])+'-'+str(dff['Città (Ubicazione Hotel)'][i])+'-'+str(dff['Cap (Ubicazione Hotel)'][i])+'- ('+str(dff['Provincia \n(Ubicazione Hotel)'][i])+')')+len('Zona rischi catastrofali Nr.: ')+1
    NS=159
    p.add_run('\nUbicazione Hotel: {}'.format(str(dff['Indirizzo (Ubicazione Hotel)'][i])+'-'+str(dff['Città (Ubicazione Hotel)'][i])+'-'+str(dff['Cap (Ubicazione Hotel)'][i])+'- ('+str(dff['Provincia \n(Ubicazione Hotel)'][i])+')')+' '*(NS-SSS))
    p.add_run('Zona rischi catastrofali Nr.: {}'.format(str(dff['Earthquake Zone'][i])))
    SSS=len('Fatturato Hotel 2021: ')+len('Fatturato Hotel stimato 2022: ')+len(dff['Fatturato  Hotel 2021 (€)'][i])+len(dff['Fatturato  Hotel 2022 (€)'][i])
    if SSS<58:
        NS=172
    elif SSS<68:
        NS=182
    elif SSS<78:
        NS=172
    elif SSS<98:
        NS=162
    elif SSS<108:
        NS=152
    print(SSS,dff['Codice_Unico'][i],22222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222)
    p.add_run('\nFatturato Hotel 2021: {}'.format(dff['Fatturato  Hotel 2021 (€)'][i])+ ' '*(NS-SSS)+ 'Fatturato Hotel stimato 2022: {}'.format(dff['Fatturato  Hotel 2022 (€)'][i]))
    p=document.add_paragraph('')#
    p.add_run('Validità della copertura').bold=True
    p.add_run('\nEffetto della copertura H 24.00 del: 30/06/2022                                                                  Scadenza della copertura H 24.00 del: 30/06/2023')
    p.add_run('\nRateazione: Annuale')
    ##
    ##
    ##
    ##
    p=document.add_paragraph()
    r=6
    table = document.add_table(rows=r, cols=2)
    table.style = 'Light Grid'
    ###################
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Responsabilità Civile'
    hdr_cells[1].text = 'Massimali Di Garanzia (combined single limit)'
    table.cell(1, 0).text='Opzione di franchigia base prescelta:'
    table.cell(1, 1).text=str(dff['Franchigia (RCT)'][i])
    table.cell(2, 0).text='Responsabilità Civile verso Terzi (RCT)'
    table.cell(2, 1).text=str(dff['RC Terzi (RCT)'][i])+' = per sinistro / anno assicurativo'
    table.cell(3, 0).text='Responsabilità Civile Prodotti (RCP) e SMERCIO'
    table.cell(3, 1).text=str(dff['RC Prodotti (RCP) e Smercio'][i])+' = per sinistro / anno assicurativo'
    table.cell(4, 0).text='Responsabilità Civile Prestatori di Lavoro (RCO)'
    table.cell(4, 1).text=str(dff['RC Prestatori di Lavoro (RCO)'][i])+' = per sinistro / anno assicurativo'
    table.cell(5, 1).text=str(dff['RC Prestatori di Lavoro (RCO) - per persona'][i])+' = per persona infortunata'
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.2)
    for nnn in np.arange(1,r):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    p=document.add_paragraph('La somma di {} si intende quale esposizione massima della compagnia anche in caso di sinistro che interessi\ncontemporaneamente la Responsabilità Civile Terzi e la Responsabilità Civile Prodotti'.format(str(dff['RC Terzi (RCT)'][i])))
    ##
    ##
    ##
    ##
    r=10
    table = document.add_table(rows=r, cols=2)
    table.style = 'Light Grid'
    ###################
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sottolimiti'
    table.cell(1, 0).text='Danni da sospensione ed interruzione d\'esercizio'
    table.cell(1, 1).text='1.000.000 € per sinistro/anno'
    table.cell(2, 0).text='Danni da incendio'
    table.cell(2, 1).text='Massimali RCT per sinistro/anno'
    table.cell(3, 0).text='Malattie professionali'
    table.cell(3, 1).text='1.500.000 € per sinistro/anno/persona'
    table.cell(4, 0).text='Cose consegnate'
    table.cell(4, 1).text='500.000 € per sinistro/anno assicurativo'
    table.cell(5, 0).text='Cose non consegnate - Garage Keeper\'s liability'
    table.cell(5, 1).text='300.000 € per sinistro/anno assicurativo'
    table.cell(6, 0).text='Servizi ai clienti'
    table.cell(6, 1).text='2.500 € per sinistro/cliente'
    table.cell(7, 0).text='Danni a beni di terzi non clienti'
    table.cell(7, 1).text='100.000 € per sinistro/anno assicurativo'
    table.cell(8, 0).text='Servizi di guardaroba e/o deposito'
    table.cell(8, 1).text='25.000 € per sinistro/anno assicurativo'
    table.cell(9, 0).text='Inquinamento accidentale (72h)'
    table.cell(9, 1).text='250.000 € per sinistro/anno assicurativo'
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.2)
    for nnn in np.arange(1,r):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    ##
    ##
    ##
    ##
    document.add_paragraph()
    r=4
    table = document.add_table(rows=r, cols=2)
    table.style = 'Light Grid'
    ###################
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Condizioni Speciali Eventi Organizzati'
    table.cell(1, 0).text='Danni patrimoniali (Financial Loss)'
    table.cell(1, 1).text=str(dff['RC Terzi (RCT) Danni Patrimoniali'][i])
    table.cell(2, 0).text='Danni alle opere d\'arte'
    table.cell(2, 1).text=str(dff['RC Terzi (RCT) Danni Opere D\'Arte'][i])
    table.cell(3, 0).text='Furto di beni consegnati ed oggetto di eventi organizzati'
    table.cell(3, 1).text=str(dff['RC Terzi (RCT) Furto di beni consegnati'][i])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.2)
    for nnn in np.arange(1,r):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    ##
    ##
    ##
    ##
    document.add_paragraph()
    r=2
    table = document.add_table(rows=r, cols=2)
    table.style = 'Light Grid'
    ###################
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Franchigia Furto RC'
    table.cell(1, 0).text='Eliminazione Franchigia Furto RC'
    table.cell(1, 1).text=str(dff['RC Terzi (RCT) Annullamento Franchigia Furto'][i])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.2)
    for nnn in np.arange(1,r):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    document.add_paragraph()
    p=document.add_paragraph('L\'Assicurato  ')
    p.add_run().add_picture('1.png', width=Inches(0.2))
    p.add_run('___________________')
    document.save('RC_{}.docx'.format(dff['Codice_Unico'][i]))
##################
##################
#BBBBBBBBBBBBBBBB
#BBBBBBBBBBBBBBBB
#BBBBBBBBBBBBBBBB
#BBBBBBBBBBBBBBBB
#BBBBBBBBBBBBBBBB
#BBBBBBBBBBBBBBBB
#BBBBBBBBBBBBBBBB
#BBBBBBBBBBBBBBBB
##################
##################

for i in range(0,len(dff['Codice Hotel'])):
    #i=1
    WW = RGBColor(255, 255, 255)
    document = Document('temp.docx')
    style = document.styles['Normal']
    font = style.font
    font.name='Times New Roman'
    ##
    ##
    ##
    ##
    SSS=len('Assicurato: ')+len(dff['Denominazione Hotel'][i])+len('Contraente: ')+len('Best Western Italia Scpa')
    p=document.add_paragraph('');
    p.add_run('Assicurato: ').bold=True;
    if SSS<53:
        NS=188
    elif SSS<58:
        NS=183
    elif SSS<63:
        NS=178
    elif SSS<68:
        NS=173
    elif SSS<73:
        NS=168
    elif SSS<78:
        NS=163
    elif SSS<83:
        NS=158
    p.add_run('{}'.format(dff['Denominazione Hotel'][i])+' '*(NS-SSS));
    p.add_run('Contraente: ').bold=True;
    p.add_run('Best Western Italia Scpa');
    SSS=len('Codice Fiscale/P.IVA: ')+len(str(dff['C.F./P.IVA'][i]))+len(str(dff['Sede Legale'][i])+'-'+str(dff['Città (Sede Legale)'][i])+'-'+str(dff['Cap\n(Sede Legale)'][i])+'- ('+str(dff['Provincia (Sede Legale)'][i])+')')+len('Sede legale: ')
    p.add_run('\nCodice Fiscale/P.IVA: ').bold=True;
    if SSS<53:
        NS=188
    elif SSS<58:
        NS=183
    elif SSS<63:
        NS=178
    elif SSS<68:
        NS=173
    elif SSS<73:
        NS=168
    elif SSS<78:
        NS=163
    elif SSS<83:
        NS=158
    elif SSS<88:
        NS=153
    elif SSS<93:
        NS=148
    elif SSS<98:
        NS=143
    p.add_run('{}'.format(dff['C.F./P.IVA'][i])+' '*(NS-SSS));
    p.add_run('Sede legale: ').bold=True;
    p.add_run('{}'.format(str(dff['Sede Legale'][i]).capitalize()+'-'+str(dff['Città (Sede Legale)'][i])+'-'+str(dff['Cap\n(Sede Legale)'][i])+'- ('+str(dff['Provincia (Sede Legale)'][i])+')'))
    p=document.add_paragraph('')#
    p.add_run('Hotel Assicurato').bold=True
    SSS=len('Denominazione Hotel: ')+len('Property Code: ')+len(dff['Denominazione Hotel'][i])+len(str(dff['Codice Hotel'][i]))
    if SSS<53:
        NS=185
    elif SSS<58:
        NS=180
    elif SSS<63:
        NS=175
    elif SSS<68:
        NS=170
    elif SSS<73:
        NS=165
    elif SSS<78:
        NS=160
    elif SSS<83:
        NS=155
    elif SSS<88:
        NS=150
    elif SSS<93:
        NS=145
    elif SSS<98:
        NS=140
    p.add_run('\nDenominazione Hotel: {}'.format(dff['Denominazione Hotel'][i])+' '*(NS-SSS)+ 'Property Code: {}'.format(dff['Codice Hotel'][i]))
    SSS=len('Ubicazione Hotel: ')+len(str(dff['Indirizzo (Ubicazione Hotel)'][i])+'-'+str(dff['Città (Ubicazione Hotel)'][i])+'-'+str(dff['Cap (Ubicazione Hotel)'][i])+'- ('+str(dff['Provincia \n(Ubicazione Hotel)'][i])+')')+len('Zona rischi catastrofali Nr.: ')+1
    NS=159
    p.add_run('\nUbicazione Hotel: {}'.format(str(dff['Indirizzo (Ubicazione Hotel)'][i])+'-'+str(dff['Città (Ubicazione Hotel)'][i])+'-'+str(dff['Cap (Ubicazione Hotel)'][i])+'- ('+str(dff['Provincia \n(Ubicazione Hotel)'][i])+')')+' '*(NS-SSS))
    p.add_run('Zona rischi catastrofali Nr.: {}'.format(str(dff['Earthquake Zone'][i])))
    SSS=len('Fatturato Hotel 2021: ')+len('Fatturato Hotel stimato 2022: ')+len(dff['Fatturato  Hotel 2021 (€)'][i])+len(dff['Fatturato  Hotel 2022 (€)'][i])
    if SSS<58:
        NS=172
    elif SSS<68:
        NS=182
    elif SSS<78:
        NS=172
    elif SSS<98:
        NS=162
    elif SSS<108:
        NS=152
    print(SSS,dff['Codice_Unico'][i],22222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222222)
    p.add_run('\nFatturato Hotel 2021: {}'.format(dff['Fatturato  Hotel 2021 (€)'][i])+ ' '*(NS-SSS)+ 'Fatturato Hotel stimato 2022: {}'.format(dff['Fatturato  Hotel 2022 (€)'][i]))
    p=document.add_paragraph('')#
    p.add_run('Validità della copertura').bold=True
    p.add_run('\nEffetto della copertura H 24.00 del: 30/06/2022                                                                  Scadenza della copertura H 24.00 del: 30/06/2023')
    p.add_run('\nRateazione: Annuale')
    ##
    ##
    ##
    ##
    p=document.add_paragraph()
    r=9
    table = document.add_table(rows=r, cols=2)
    table.style = 'Light Grid'
    ###################
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Danni diretti/Partite Assicurate'
    hdr_cells[1].text = 'Somme Assicurate'
    table.cell(1, 0).text='Opzione Franchigia Frontale Prescelta';
    table.cell(1, 1).text=str(dff['Franchigia Danni Diretti'][i])
    table.cell(2, 0).text='1. Fabbricato'
    table.cell(2, 1).text=str(dff['Fabbricato'][i])
    table.cell(3, 0).text='2. Contenuto'
    table.cell(3, 1).text=str(dff['Contenuto'][i])
    table.cell(4, 0).text='3. Ricorso Terzi'
    table.cell(4, 1).text=str(dff['Ricorso Terzi'][i])
    table.cell(5, 0).text='4. Cristalli/Insegne'
    table.cell(5, 1).text=str(dff['Cristalli'][i])
    table.cell(6, 0).text='5. Furto Contenuto'
    table.cell(6, 1).text=str(dff['Furto'][i])
    table.cell(7, 0).text='6. Fenomeno Elettrico'
    table.cell(7, 1).text=str(dff['Fenomeno Elettrico'][i])
    table.cell(8, 0).text='7. Merci In Refrigerazione'
    table.cell(8, 1).text=str(dff['Merci in Refrigerazione'][i])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.2)
    for nnn in np.arange(1,r):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    ##
    ##
    ##
    ##
    document.add_paragraph()
    r=6
    table = document.add_table(rows=r, cols=2)
    table.style = 'Light Grid'
    ###################
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Danni Indiretti'
    table.cell(1, 0).text='Operatività'
    if len(dff['Margine di contribuzione annuo'][i].strip())>3:
        table.cell(1, 1).text='Operante'
    else:
        table.cell(1, 1).text='Non Operante'
    table.cell(2, 0).text='Margine di contribuzione annuo'
    table.cell(2, 1).text=str(dff['Margine di contribuzione annuo'][i])
    table.cell(3, 0).text='- Periodo di Indennizzo (Mesi)'
    table.cell(3, 1).text=str(dff['Periodo di Indennizzo (Mesi)'][i])
    table.cell(4, 0).text='Danni da cimici da letto, legionella e salmonella'
    table.cell(4, 1).text=str(dff['Cimici da letto'][i])
    table.cell(5, 0).text='- Periodo di Indennizzo (Mesi)'
    table.cell(5, 1).text=str(dff['Cimici da letto\n Periodo di Indennizzo (n.mesi)'][i])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.2)
    for nnn in np.arange(1,r):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    ##
    ##
    ##
    ##
    document.add_paragraph()
    r=3
    table = document.add_table(rows=r, cols=2)
    table.style = 'Light Grid'
    ###################
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Garanzie facoltative (valide per le sezioni operanti)'
    table.cell(1, 0).text='Terremoto, Inondazione, Alluvione, Allagamento'
    table.cell(1, 1).text=str(dff['Terremoto, Inondazione, Alluvione, Allagamento'][i])
    table.cell(2, 0).text='Terrorismo, Eventi Socio Politici, Atti Dolosi'
    table.cell(2, 1).text=str(dff['Terrorismo, Eventi Socio Politici, Atti Dolosi'][i])
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.2)
    for nnn in np.arange(1,r):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    ##
    ##
    ##
    ##
    document.add_paragraph()
    r=2
    table = document.add_table(rows=r, cols=4)
    table.style = 'Table Grid'
    ###################
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Conteggio del Premio'
    hdr_cells[1].text = 'Imponibile'
    hdr_cells[2].text = 'Imposte'
    hdr_cells[3].text = 'Lordo'
    table.cell(1, 0).text='alla firma'
    #
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height=Inches(0.2)
    for nnn in np.arange(1,r):
        table.rows[nnn].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[nnn].height=Inches(0.2)
    #
    document.add_paragraph()
    p=document.add_paragraph('L\'Assicurato  ')
    p.add_run().add_picture('1.png', width=Inches(0.2))
    p.add_run('___________________')
    document.save('Fab_{}.docx'.format(dff['Codice_Unico'][i]))
